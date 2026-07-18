# -*- coding: utf-8 -*-
"""Generate Spanish and English progress report .docx files."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]

INK = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x6B, 0x72, 0x80)
ACCENT = RGBColor(0xE8, 0x70, 0x0A)

LOCALES = {
    "es": {
        "source": ROOT / "REPORTES_DIARIOS.md",
        "output": ROOT / "reporte.docx",
        "start_line": 2582,
        "end_line": 3255,
        "report_date": "17 de julio de 2026",
        "subtitle": "Reporte de avance",
        "task_prefix": "Tarea",
        "intro": (
            "Resumen de implementación de la aplicación móvil para conductores: "
            "autenticación, listado de cargas, progreso en ruta, equipo requerido "
            "y experiencia de uso."
        ),
        "footer": "Tigerhawk Mobile · Recruiting Smarter Brasil",
        "task_heading": r"^### Tarea \d+ — ",
        "section_labels": {
            "implemented": "Qué se implementó",
            "available": "Funcionalidad disponible",
            "testing": "Cómo probar",
        },
        "skip_prefixes": (
            "npm ",
            "abrir ",
            "confirmar en ",
            "revisar ",
            "diff clave",
            "_al cerrar",
            "open ",
            "confirm in ",
            "review ",
            "key diff",
            "_when closing",
        ),
        "title_replacements": {
            r"\(`z-feedback_cliente`\)$": "",
            r"\(`TASKS\.md` \+ docs\)$": "",
            r"Alinear TASKS\.md a RESPUESTAS_CLIENTE": (
                "Alinear backlog a respuestas del cliente"
            ),
            r"Auditoría TMS_fusion \+ bridge email A\.2": (
                "Auditoría del TMS y respaldo de login por email"
            ),
            r"PR Mobile API \+ correcciones de CI en la rama del cliente": (
                "Integración de la API móvil en el TMS"
            ),
            r"Merge a main y confirmación en Netlify": (
                "Despliegue del TMS en el entorno de desarrollo"
            ),
        },
    },
    "en": {
        "source": ROOT / "DAILY_REPORTS.md",
        "output": ROOT / "reporte-en.docx",
        "start_line": 2424,
        "end_line": 3094,
        "report_date": "July 17, 2026",
        "subtitle": "Progress Report",
        "task_prefix": "Task",
        "intro": (
            "Summary of Tigerhawk Mobile driver app delivery: authentication, "
            "load lists, route progress, required equipment, and usability."
        ),
        "footer": "Tigerhawk Mobile · Recruiting Smarter Brasil",
        "task_heading": r"^### Task \d+ — ",
        "section_labels": {
            "implemented": "What was implemented",
            "available": "Functionality available",
            "testing": "How to test",
        },
        "skip_prefixes": (
            "npm ",
            "open ",
            "confirm in ",
            "review ",
            "key diff",
            "_when closing",
            "abrir ",
            "confirmar en ",
            "revisar ",
        ),
        "title_replacements": {
            r"\(`z-feedback_cliente`\)$": "",
            r"\(`TASKS\.md` \+ docs\)$": "",
            r"Align TASKS\.md to client answers": "Align backlog to client answers",
            r"TMS_fusion audit \+ A\.2 email bridge": (
                "TMS audit and email login fallback"
            ),
            r"Mobile API PR \+ CI fixes on the client branch": (
                "Mobile API integration in the TMS"
            ),
            r"Merge to main and Netlify confirmation": (
                "TMS deployment to the development environment"
            ),
        },
    },
}

FILE_PATTERNS = [
    r"`[^`]+`",
    r"TASKS\.md[^\n]*",
    r"z-feedback_cliente[^\n]*",
    r"docs/[^\n]*",
    r"supabase/[^\n]*",
    r"TMS_fusion[^\n]*",
    r"npm (run |test)[^\n]*",
    r"(Abrir|Open) [^\n]+",
    r"(Confirmar en|Confirm in) [^\n]+",
    r"(Revisar|Review) [^\n]+",
    r"(Diff clave|Key diff):[^\n]*",
    r"PR #\d+[^\n]*",
    r"commit `[^`]+`[^\n]*",
    r"(En|In) `[^`]+`:[^\n]*",
    r"(desde|from) `[^`]+`[^\n]*",
    r"— probe [^\n]+",
    r"§\d+[^\n]*",
    r"\(commit [^)]+\)",
    r"@ `main`[^\n]*",
    r"host [^\n]*netlify[^\n]*",
    r"EXPO_PUBLIC_[^\n]*",
    r"\b\d{1,2} Jul(y)?\b",
    r"July-\d+",
    r"RESPUESTAS_CLIENTE\.md",
    r"ANALISIS\.md",
    r"ANALYSIS\.docx",
    r"ANALISIS\.docx",
]


def strip_path(value: str) -> str:
    if any(
        token in value
        for token in ("/", "\\", ".md", ".tsx", ".ts", ".sql", ".mjs", "npm ")
    ):
        return ""
    return value


def clean(text: str, title_replacements: dict[str, str]) -> str:
    if not text or not text.strip():
        return ""
    t = text.strip()
    t = re.sub(r"\*\*([^*]+)\*\*", r"\1", t)
    t = re.sub(r"`([^`]+)`", lambda m: strip_path(m.group(1)), t)
    for pat in FILE_PATTERNS:
        t = re.sub(pat, "", t, flags=re.IGNORECASE)
    t = re.sub(r"→ \[[x~\-!]\]", "", t)
    t = re.sub(r"\s*→\s*", ": ", t)
    for pattern, replacement in title_replacements.items():
        t = re.sub(pattern, replacement, t)
    t = re.sub(r"\s{2,}", " ", t)
    t = re.sub(r"\s+([,.;])", r"\1", t)
    return t.strip(" ·-—").strip()


def is_skip_line(line: str, skip_prefixes: tuple[str, ...]) -> bool:
    low = line.lower().strip()
    if not low:
        return True
    if low.startswith(skip_prefixes):
        return True
    if "testpathpattern" in low:
        return True
    if low == "---":
        return True
    return False


def parse_tasks(md: str, locale: dict) -> list[dict]:
    tasks: list[dict] = []
    current: dict | None = None
    section: str | None = None
    labels = locale["section_labels"]
    label_to_key = {v: k for k, v in labels.items()}

    for line in md.splitlines():
        if line.startswith("## "):
            continue
        if re.match(locale["task_heading"], line):
            if current:
                tasks.append(current)
            title = clean(re.sub(locale["task_heading"], "", line).strip(), locale["title_replacements"])
            current = {"title": title, "sections": {}}
            section = None
            continue
        if line.startswith("**") and line.endswith("**"):
            label = line.strip("*").strip()
            if label in label_to_key:
                section = label_to_key[label]
                if current is not None:
                    current["sections"][section] = []
            continue
        if section and current is not None:
            if line.startswith("- "):
                item = clean(line[2:], locale["title_replacements"])
            elif re.match(r"^\d+\. ", line):
                item = clean(re.sub(r"^\d+\. ", "", line), locale["title_replacements"])
            else:
                continue
            if item and not is_skip_line(item, locale["skip_prefixes"]):
                current["sections"][section].append(item)

    if current:
        tasks.append(current)
    return tasks


def add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "E5E7EB")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def style_runs(paragraph, size=11, bold=False, color=INK, space_after=6):
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = color
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.15


def add_label(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = ACCENT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        style_runs(p, size=10.5, color=INK, space_after=4)


def build_document(locale: dict) -> Document:
    source: Path = locale["source"]
    markdown = "\n".join(
        source.read_text(encoding="utf-8").splitlines()[
            locale["start_line"] - 1 : locale["end_line"]
        ]
    )
    tasks = parse_tasks(markdown, locale)
    labels = locale["section_labels"]

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Tigerhawk Mobile")
    run.font.name = "Calibri"
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = INK
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(locale["subtitle"])
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(2)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(locale["report_date"])
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = ACCENT
    date_p.paragraph_format.space_after = Pt(18)

    add_horizontal_rule(doc)

    intro = doc.add_paragraph(locale["intro"])
    style_runs(intro, size=10.5, color=MUTED, space_after=14)

    for idx, task in enumerate(tasks, start=1):
        h = doc.add_paragraph()
        num_run = h.add_run(f"{locale['task_prefix']} {idx}  ")
        num_run.font.name = "Calibri"
        num_run.font.size = Pt(13)
        num_run.font.bold = True
        num_run.font.color.rgb = ACCENT
        title_run = h.add_run(task["title"])
        title_run.font.name = "Calibri"
        title_run.font.size = Pt(13)
        title_run.font.bold = True
        title_run.font.color.rgb = INK
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)

        for key in ("implemented", "available", "testing"):
            items = task["sections"].get(key, [])
            if not items:
                continue
            add_label(doc, labels[key])
            add_bullets(doc, items)

        if idx < len(tasks):
            add_horizontal_rule(doc)

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(locale["footer"])
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED

    output: Path = locale["output"]
    doc.save(output)
    print(f"Created {output} ({len(tasks)} tasks)")
    return doc


def main() -> None:
    targets = sys.argv[1:] if len(sys.argv) > 1 else list(LOCALES.keys())
    for key in targets:
        if key not in LOCALES:
            raise SystemExit(f"Unknown locale: {key}. Use: es, en")
        build_document(LOCALES[key])


if __name__ == "__main__":
    main()
