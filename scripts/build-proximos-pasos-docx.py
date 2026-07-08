"""Generate PROXIMOS_PASOS.docx (EN) and PROXIMOS_PASOS_ES.docx (ES) — July 2, 2026."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
OUT_EN = ROOT / "PROXIMOS_PASOS.docx"
OUT_ES = ROOT / "PROXIMOS_PASOS_ES.docx"


def set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def add_h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def add_h3(doc: Document, text: str) -> None:
    doc.add_heading(text, level=3)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def init_doc() -> Document:
    doc = Document()
    set_normal_style(doc)
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    return doc


def build_en() -> None:
    doc = init_doc()

    add_title(doc, "Next Steps — Tigerhawk Mobile (PP2)")
    add_para(doc, "Report date: July 2, 2026")
    doc.add_paragraph()

    add_para(
        doc,
        "Context: backlog items 9.1–9.7 are complete (transitions, documents, offline queue, "
        "Samsara live-ready integration, GPS sign-off, client handoff). The product is in "
        "field validation, deployment, and client operations phase.",
    )
    add_para(
        doc,
        "Technical detail: PP2_TAREAS_DEV.md · Handoff: docs/CLIENT_HANDOFF_9_7.md · "
        "Post-v1.1: docs/BACKLOG_V1_1_7_7.md",
    )

    add_h2(doc, "1. Immediate priorities (next two weeks)")
    add_table(
        doc,
        ["#", "Action", "Owner", "Notes"],
        [
            [
                "1",
                "Access: GitHub (mobile + TMS), Expo/EAS, Vercel (collaborator)",
                "Client",
                "Admin on main not required; merges via PR",
            ],
            [
                "2",
                "Deploy TMS dev to Netlify with latest changes",
                "Dev",
                "Includes GPS map 8.12–8.13, wait time, POD auto-stop WT.28, Samsara webhook 9.5",
            ],
            [
                "3",
                "Updated Android APK on EAS (npm run build:android:preview)",
                "Dev",
                "Same EXPO_PUBLIC_TMS_API_URL as Netlify",
            ],
            [
                "4",
                "Manual GPS sign-off G1 (docs/GPS_LIVE_TRACKING_SIGNOFF_8_17.md)",
                "QA + Client",
                "See field test section 3",
            ],
            [
                "5",
                "Sign-off: POD closes wait timer",
                "QA + Client",
                "Check In → upload POD → timer Stopped; requires TMS WT.28 deployed",
            ],
            [
                "6",
                "Samsara live (optional)",
                "Client + Dev",
                "Netlify env: SAMSARA_ENABLED, token, webhook; does not block mobile GPS",
            ],
        ],
    )

    add_h2(doc, "2. Live GPS vs Samsara — client clarification")
    add_table(
        doc,
        ["Question", "Answer"],
        [
            ["Does the app POST location to Samsara?", "No."],
            [
                "How does dispatch see the driver on the map?",
                "Phone → Supabase (loads.current_*) → TMS (blue Driver marker).",
            ],
            [
                "What does Samsara do?",
                "Geofence webhook on TMS → auto-closes wait time when leaving the site. "
                "It does not replace mobile GPS.",
            ],
            [
                "Does EDI report driver location?",
                "Not for live map; EDI is a separate business/TMS flow.",
            ],
        ],
    )

    add_h2(doc, "3. Field test (recommended) — mobile GPS, not Samsara")
    add_para(
        doc,
        "A real-world test is recommended: one person walks with the app on a phone while "
        "another person watches the TMS map for movement.",
    )
    add_para(
        doc,
        "Important: this test validates mobile GPS + Supabase + TMS map, not Samsara. "
        "Samsara requires a separate scenario (vehicle + geofence + open wait timer).",
        bold=True,
    )

    add_h3(doc, "Can the development TMS do this?")
    add_para(
        doc,
        "Yes — the code already exists in the editable TMS repository "
        "(proyecto_1_TigerHawk TMS\\tigerhawk-tms-main\\tigerhawk-tms-main).",
    )
    add_table(
        doc,
        ["Component", "File"],
        [
            ["Driver marker on map", "components/maps/LoadSidebarMap.tsx"],
            ["Realtime GPS on detail", "hooks/useLoadLiveLocation.ts"],
            ["Load detail panel", "components/dispatcher/LoadDetailPanel.tsx"],
            ["Board column", "components/dispatcher/LoadsTable.tsx (Driver Last Seen)"],
        ],
    )
    add_para(
        doc,
        "Requirement: these files must be deployed on Netlify. Without tasks 8.12–8.13, "
        "the map shows static stops only (pickup/delivery), not a moving blue Driver point.",
    )

    add_h3(doc, "Field test checklist (~30 minutes)")
    add_para(doc, "Before", bold=True)
    add_bullets(
        doc,
        [
            "Supabase SQL applied: 20260605120000_pp2_driver_live_location_loads.sql, "
            "fix_pp2_driver_location_trigger_updated_at.sql, enable_realtime_driver_tracking.sql",
            "TMS Netlify deployed with package 8.12–8.13 (see docs/TMS_DEV_REPOSITORY.md)",
            "APK with same Supabase and TMS API URL as production project",
            "Test driver assigned to one active load (e.g. In Transit)",
            "Physical phone (not mobile web — GPS banner does not run on web)",
        ],
    )
    add_para(doc, "During", bold=True)
    add_numbered(
        doc,
        [
            "Driver: login → My Loads → open load → allow location → keep detail screen open",
            "Dispatcher: TMS → same load → side panel → map with Driver legend",
            "Driver walks 50–100 m (or simulate location on emulator)",
            "Wait ≤ 60 s between updates (ping ~45 s)",
        ],
    )
    add_para(doc, "Success criteria", bold=True)
    add_table(
        doc,
        ["Surface", "Expected result"],
        [
            ["Mobile", "Banner: Last sent: Just now"],
            ["TMS board", "Driver Last Seen is not empty"],
            ["TMS detail map", "Blue Driver marker updates when moving"],
            ["Supabase", "loads.current_latitude, current_longitude, last_seen_at are recent"],
        ],
    )
    add_para(
        doc,
        "If the test fails: apply fix_pp2_driver_location_trigger_updated_at.sql; verify Realtime; "
        "confirm Netlify has useLoadLiveLocation deployed.",
    )
    add_para(doc, "Full matrix: docs/QA_DRIVER_LIVE_TRACKING.md (cases G1–G9).")

    add_h2(doc, "4. Samsara test (separate, when credentials are available)")
    add_para(doc, "Does not replace the field test in section 3.")
    add_numbered(
        doc,
        [
            "Mobile: Check In (open wait at delivery)",
            "Samsara: delivery geofence exit or POST …/api/integrations/samsara/simulate (staff)",
            "TMS: wait Stopped + activity_log delivery_wait_geofence_auto_stop",
        ],
    )
    add_para(doc, "Guide: docs/SAMSARA_GEOFENCE_SPIKE.md · docs/QA_SAMSARA_GEOFENCE_MOCK.md")

    add_h2(doc, "5. Git and deployment process")
    add_table(
        doc,
        ["Repository", "Purpose", "Deploy"],
        [
            ["proyecto_PP2_app_mobile", "Mobile development", "EAS → APK"],
            ["tigerhawk-tms-main", "TMS development", "Vercel/Netlify"],
            ["tigerhawk_mobile (client GitHub)", "Stable mirror", "npm run sync:client-mirror"],
        ],
    )
    add_para(
        doc,
        "Routine: feature branch → PR + CI → staging review → merge to main → "
        "automatic TMS redeploy / EAS build when releasing.",
    )
    add_para(
        doc,
        "The client coordinator does not need to manage branches daily: approve merges "
        "and validate on staging.",
    )

    add_h2(doc, "6. After field validation")
    add_table(
        doc,
        ["Topic", "Reference"],
        [
            ["Push / messages / E2E", "docs/BACKLOG_V1_1_7_7.md"],
            ["iOS TestFlight", "Apple Developer account + EAS"],
            ["Formal client delivery", "docs/CLIENT_HANDOFF_9_7.md §9 checklist"],
            ["Operations support", "docs/MOBILE_SUPPORT_RUNBOOK_7_6.md"],
        ],
    )

    add_h2(doc, "7. Summary status (July 2026)")
    add_table(
        doc,
        ["Area", "Status"],
        [
            ["Login, loads, actions, documents, offline", "Complete"],
            ["Manual wait time + detention emails", "Complete"],
            ["Driver POD closes timer (WT.28)", "Code complete · pending TMS deploy validation"],
            ["Live GPS mobile → TMS", "Code complete · field test recommended"],
            ["Samsara geofence", "Ready in TMS dev · pending client credentials + webhook"],
            ["Production APK / iOS", "Pending access + EAS build"],
        ],
    )
    add_para(
        doc,
        "Suggested next milestone: 30-minute meeting — access setup + one field test (§3) "
        "+ one POD/wait flow (§1 item 5).",
    )

    add_h1(doc, "Deploying the iOS app with an Apple Developer Account")
    add_para(
        doc,
        "To publish or distribute the app on iPhone/iPad (TestFlight or App Store), Apple "
        "requires an active developer account. This is Apple’s official program (~USD 99/year), "
        "not a fee charged by the development team. It enables app signing, upload to Apple "
        "servers, and secure installation on iOS devices.",
    )
    add_para(doc, "Why is this license required?", bold=True)
    add_para(
        doc,
        "Apple does not allow production apps on iPhones without its signing and review process. "
        "The membership covers certificates, provisioning profiles, and access to TestFlight / "
        "App Store Connect. Without it, testing is limited to the simulator or strict development modes.",
    )
    add_para(doc, "Whose name appears on the app?", bold=True)
    add_table(
        doc,
        ["Option", "What the iPhone user sees", "Who pays Apple"],
        [
            [
                "Publish with the development team’s account",
                "The app lists the development team’s organization on App Store / TestFlight",
                "Development team (or as agreed in contract)",
            ],
            [
                "Publish with the client’s account",
                "The app lists the client’s company name and signature",
                "Client (Apple license in their name)",
            ],
        ],
    )
    add_para(
        doc,
        "Both options are valid — this is a branding and ownership decision, not a quality issue. "
        "If the client wants their company shown as the publisher, the Apple license should be "
        "in the client’s name. If speed is the priority, the development team can publish using "
        "its own account; App Store/TestFlight would list the vendor as developer until migration "
        "to the client account.",
    )
    add_para(doc, "How the client can proceed (when ready for iOS):", bold=True)
    add_numbered(
        doc,
        [
            "On their own — Create an account at https://developer.apple.com, pay the annual fee, and "
            "publish (the development team can guide EAS/Expo steps if needed).",
            "Delegated to the development team — The client pays the license in their name and "
            "securely shares App Store Connect / signing credentials; the team builds and "
            "uploads to TestFlight or App Store under the client account.",
        ],
    )
    add_para(
        doc,
        "Android does not require an equivalent step for an internal APK. iOS does when "
        "running real pilots on iPhones outside the simulator.",
    )
    add_para(doc, "Technical detail: docs/MOBILE_BUILDS.md §5 · docs/EAS_CREDENTIALS_HANDOFF_7_5.md")

    add_h1(doc, "What are GitHub repository invitations for?")
    add_para(doc, "Invitations to these two repositories:")
    add_bullets(
        doc,
        [
            "TMS Tigerhawk V2 — TMS (web panel / Vercel server): "
            "https://github.com/AI-Developer-Smarter-Business/TMS_tigerhawk_V2",
            "Tigerhawk Mobile App — mobile app (Expo): "
            "https://github.com/AI-Developer-Smarter-Business/tigerhawk_mobile_app",
        ],
    )
    add_para(
        doc,
        "They are not meant to force everyone to program daily. They ensure the client and "
        "the development team work from the same copy of the project.",
    )
    add_para(
        doc,
        "Think of it as a shared cloud folder with version history:",
    )
    add_bullets(
        doc,
        [
            "When the development team pushes a fix or feature, the client can pull it and have the same code.",
            "When the client (or an AI assistant on their machine) pushes a change, the team sees it and continues from there.",
            "This avoids two divergent versions (developer vs office) that no longer match production.",
        ],
    )
    add_para(
        doc,
        "This prevents issues such as: “it works in my environment but not on deployed TMS” "
        "or “the driver app does not have what we tested yesterday.”",
    )
    add_para(doc, "In practice:", bold=True)
    add_numbered(
        doc,
        [
            "Accept the invitation email on GitHub.",
            "Clone or update the repository when the latest code is needed.",
            "Before making changes: pull the latest code, then work, then push (if the client publishes changes).",
        ],
    )
    p = doc.add_paragraph()
    run = p.add_run(
        "Important note: Before pushing changes to the repository (git push), always run "
        "git pull first. This downloads the latest code from GitHub and merges it with local "
        "work. Pushing without pulling first can overwrite others’ work or create hard-to-fix "
        "conflicts — like editing an outdated document.",
    )
    run.bold = True
    add_para(
        doc,
        "In short: GitHub invitations keep TMS and the mobile app in one source of truth; "
        "client and developer see the same files and history, reducing compatibility issues.",
    )
    add_para(doc, "Brief guide: docs/GitHub_Setup_Guide.md")

    doc.save(OUT_EN)
    print(f"Wrote {OUT_EN}")


def build_es() -> None:
    doc = init_doc()

    add_title(doc, "Próximos pasos — Tigerhawk Mobile (PP2)")
    add_para(doc, "Fecha del reporte: 2 de julio de 2026")
    doc.add_paragraph()

    add_para(
        doc,
        "Contexto: backlog 9.1–9.7 cerrado (transiciones, documentos, offline, integración "
        "Samsara live-ready, sign-off GPS, handoff). El producto está en fase de validación en "
        "campo, deploy y operación con el cliente.",
    )
    add_para(
        doc,
        "Detalle técnico: PP2_TAREAS_DEV.md · Handoff: docs/CLIENT_HANDOFF_9_7.md · "
        "Post-v1.1: docs/BACKLOG_V1_1_7_7.md",
    )

    add_h2(doc, "1. Prioridad inmediata (esta quincena)")
    add_table(
        doc,
        ["#", "Acción", "Responsable", "Notas"],
        [
            [
                "1",
                "Accesos: GitHub (móvil + TMS), Expo/EAS, Vercel (collaborator)",
                "Cliente",
                "No hace falta admin de main; merges vía PR",
            ],
            [
                "2",
                "Desplegar TMS dev en Netlify con últimos cambios",
                "Dev",
                "Incluye mapa GPS 8.12–8.13, wait time, POD auto-stop WT.28, webhook Samsara 9.5",
            ],
            [
                "3",
                "APK Android actualizado en EAS (npm run build:android:preview)",
                "Dev",
                "Mismo EXPO_PUBLIC_TMS_API_URL que Netlify",
            ],
            [
                "4",
                "Sign-off manual GPS G1 (docs/GPS_LIVE_TRACKING_SIGNOFF_8_17.md)",
                "QA + cliente",
                "Ver prueba en calle §3",
            ],
            [
                "5",
                "Sign-off: POD cierra wait timer",
                "QA + cliente",
                "Check In → subir POD → timer Stopped; requiere TMS WT.28 desplegado",
            ],
            [
                "6",
                "Samsara live (opcional)",
                "Cliente + Dev",
                "Env Netlify: SAMSARA_ENABLED, token, webhook; no bloquea GPS móvil",
            ],
        ],
    )

    add_h2(doc, "2. GPS en vivo vs Samsara — aclaración para el cliente")
    add_table(
        doc,
        ["Pregunta", "Respuesta"],
        [
            ["¿La app hace POST a Samsara para ubicación?", "No."],
            [
                "¿Cómo ve dispatch al conductor en el mapa?",
                "Teléfono → Supabase (loads.current_*) → TMS (marcador azul Driver).",
            ],
            [
                "¿Qué hace Samsara?",
                "Webhook geofence en el TMS → auto cierre de wait time al salir del sitio. "
                "No sustituye el GPS del móvil.",
            ],
            [
                "¿EDI reporta ubicación del conductor?",
                "No para el mapa en vivo; EDI es otro flujo de negocio/TMS.",
            ],
        ],
    )

    add_h2(doc, "3. Prueba en calle (recomendada) — GPS móvil, no Samsara")
    add_para(
        doc,
        "Sí, es buena idea hacer una prueba real: una persona con la app en el teléfono camina "
        "y otra persona en el TMS observa si el punto se mueve.",
    )
    add_para(
        doc,
        "Importante: esa prueba valida GPS del móvil + Supabase + mapa TMS, no Samsara. "
        "Para Samsara haría falta otro escenario (camión con geofence + wait abierto).",
        bold=True,
    )

    add_h3(doc, "¿Puede hacerlo el TMS en desarrollo?")
    add_para(
        doc,
        "Sí, el código ya está en el repo TMS editable "
        "(proyecto_1_TigerHawk TMS\\tigerhawk-tms-main\\tigerhawk-tms-main).",
    )
    add_table(
        doc,
        ["Pieza", "Archivo / componente"],
        [
            ["Marcador conductor en mapa", "components/maps/LoadSidebarMap.tsx"],
            ["Realtime GPS en detalle", "hooks/useLoadLiveLocation.ts"],
            ["Panel detalle carga", "components/dispatcher/LoadDetailPanel.tsx"],
            ["Columna board", "components/dispatcher/LoadsTable.tsx (Driver Last Seen)"],
        ],
    )
    add_para(
        doc,
        "Condición: esos archivos deben estar desplegados en Netlify. Sin las tareas 8.12–8.13, "
        "el mapa solo mostrará paradas estáticas (pickup/delivery), sin punto azul en movimiento.",
    )

    add_h3(doc, "Checklist prueba en calle (≈ 30 min)")
    add_para(doc, "Antes", bold=True)
    add_bullets(
        doc,
        [
            "SQL Supabase aplicado: 20260605120000_pp2_driver_live_location_loads.sql, "
            "fix_pp2_driver_location_trigger_updated_at.sql, enable_realtime_driver_tracking.sql",
            "TMS Netlify con paquete 8.12–8.13 (ver docs/TMS_DEV_REPOSITORY.md)",
            "APK con EXPO_PUBLIC_SUPABASE_* y EXPO_PUBLIC_TMS_API_URL del mismo proyecto",
            "Conductor de prueba asignado a 1 carga en estado activo (p. ej. In Transit)",
            "Teléfono físico (no navegador web — el banner GPS no corre en web)",
        ],
    )
    add_para(doc, "Durante", bold=True)
    add_numbered(
        doc,
        [
            "Conductor: login → My Loads → abrir la carga → permitir ubicación → mantener el detalle abierto",
            "Dispatcher: TMS → misma carga → panel lateral → mapa con leyenda Driver",
            "Conductor camina 50–100 m (o simula ubicación en emulador)",
            "Esperar ≤ 60 s entre actualizaciones (ping ~45 s)",
        ],
    )
    add_para(doc, "Éxito", bold=True)
    add_table(
        doc,
        ["Dónde", "Qué debe verse"],
        [
            ["Móvil", "Banner: Last sent: Just now"],
            ["TMS board", "Driver Last Seen ≠ —"],
            ["TMS mapa detalle", "Marcador azul Driver que se actualiza al moverse"],
            ["Supabase", "loads.current_latitude, current_longitude, last_seen_at recientes"],
        ],
    )
    add_para(
        doc,
        "Si falla: aplicar fix_pp2_driver_location_trigger_updated_at.sql; verificar Realtime; "
        "confirmar que Netlify tiene useLoadLiveLocation desplegado.",
    )
    add_para(doc, "Matriz completa: docs/QA_DRIVER_LIVE_TRACKING.md (casos G1–G9).")

    add_h2(doc, "4. Prueba Samsara (separada, cuando haya credenciales)")
    add_para(doc, "No sustituye la prueba en calle del §3.")
    add_numbered(
        doc,
        [
            "Móvil: Check In (wait abierto en entrega)",
            "Samsara: salida de geofence de entrega o POST …/api/integrations/samsara/simulate (staff)",
            "TMS: wait Stopped + activity_log delivery_wait_geofence_auto_stop",
        ],
    )
    add_para(doc, "Guía: docs/SAMSARA_GEOFENCE_SPIKE.md · docs/QA_SAMSARA_GEOFENCE_MOCK.md")

    add_h2(doc, "5. Proceso Git / deploy (para el cliente)")
    add_table(
        doc,
        ["Repo", "Uso", "Deploy"],
        [
            ["proyecto_PP2_app_mobile", "Desarrollo app", "EAS → APK"],
            ["tigerhawk-tms-main", "Desarrollo TMS", "Vercel/Netlify"],
            ["tigerhawk_mobile (GitHub cliente)", "Mirror estable", "npm run sync:client-mirror"],
        ],
    )
    add_para(
        doc,
        "Rutina: rama feature/… → PR + CI → revisión en staging → merge main → "
        "redeploy automático TMS / build EAS cuando toque release.",
    )
    add_para(
        doc,
        "Quien coordina del lado del cliente no necesita mantener ramas a diario: "
        "aprueba merges y valida en staging.",
    )

    add_h2(doc, "6. Después de validación en campo")
    add_table(
        doc,
        ["Tema", "Referencia"],
        [
            ["Push / mensajes / E2E", "docs/BACKLOG_V1_1_7_7.md"],
            ["iOS TestFlight", "Cuenta Apple Developer + EAS"],
            ["Entrega formal cliente", "docs/CLIENT_HANDOFF_9_7.md §9 checklist"],
            ["Soporte operativo", "docs/MOBILE_SUPPORT_RUNBOOK_7_6.md"],
        ],
    )

    add_h2(doc, "7. Estado resumido (jul 2026)")
    add_table(
        doc,
        ["Área", "Estado"],
        [
            ["Login, cargas, acciones, documentos, offline", "Completo"],
            ["Wait time manual + emails detention", "Completo"],
            ["POD conductor cierra timer (WT.28)", "Código listo · pendiente validar con TMS desplegado"],
            ["GPS en vivo móvil → TMS", "Código listo · prueba en calle recomendada"],
            ["Samsara geofence", "Listo en TMS dev · pendiente credenciales + webhook cliente"],
            ["APK producción / iOS", "Pendiente accesos + build EAS"],
        ],
    )
    add_para(
        doc,
        "Siguiente hito sugerido: reunión 30 min — accesos + 1 prueba en calle (§3) "
        "+ 1 flujo POD/wait (§1 fila 5).",
    )

    add_h1(doc, "Desplegar aplicación para iOS con Apple Developer Account")
    add_para(
        doc,
        "Para publicar o distribuir la app en iPhone/iPad (TestFlight o App Store), Apple exige "
        "una cuenta de desarrollador activa. No es un cobro del equipo de desarrollo: es el "
        "programa oficial de Apple (~99 USD/año) que permite firmar la app, subirla a sus "
        "servidores y que los conductores la instalen de forma segura en dispositivos iOS.",
    )
    add_para(doc, "¿Por qué hay que pagar esa licencia?", bold=True)
    add_para(
        doc,
        "Apple no permite instalar apps de producción en iPhones sin pasar por su proceso de "
        "firma y revisión. La membresía cubre certificados, perfiles de instalación y el acceso "
        "a TestFlight/App Store Connect. Sin ella, solo se puede probar en simulador o con "
        "limitaciones muy estrictas en desarrollo.",
    )
    add_para(doc, "¿A nombre de quién sale la app?", bold=True)
    add_table(
        doc,
        ["Opción", "Qué ve el usuario en el iPhone", "Quién paga Apple"],
        [
            [
                "Publicar con la cuenta del equipo de desarrollo",
                "La app aparece bajo el nombre de la organización del equipo en App Store / TestFlight",
                "Equipo de desarrollo (o según contrato)",
            ],
            [
                "Publicar con la cuenta del cliente",
                "La app aparece bajo el nombre y la firma de su empresa",
                "Cliente (licencia Apple a su nombre)",
            ],
        ],
    )
    add_para(
        doc,
        "No hay problema con ninguna de las dos vías: es una decisión de marca y propiedad, "
        "no de calidad técnica. Si el cliente prefiere que en el teléfono figure su empresa "
        "como publicador, lo más natural es que la licencia Apple esté a su nombre. Si en "
        "este momento prefieren avanzar rápido y el equipo de desarrollo publica con su propia "
        "cuenta, también se puede; solo quedaría claro que en App Store/TestFlight el "
        "desarrollador listado sería el del proveedor hasta migrar a la cuenta del cliente.",
    )
    add_para(doc, "Cómo puede hacerlo el cliente (cuando decidan iOS):", bold=True)
    add_numbered(
        doc,
        [
            "Por su cuenta — Crean la cuenta en https://developer.apple.com, pagan la anualidad "
            "y publican (el equipo de desarrollo guía el proceso técnico con EAS/Expo si hace falta).",
            "Delegado al equipo de desarrollo — El cliente paga la licencia a su nombre y comparte "
            "acceso a App Store Connect / credenciales de firma (de forma segura); el equipo "
            "genera el build y lo sube a TestFlight o App Store en la cuenta del cliente.",
        ],
    )
    add_para(
        doc,
        "Android no requiere este paso equivalente para un APK interno; iOS sí cuando quieran "
        "pilotos reales en iPhone fuera del simulador.",
    )
    add_para(doc, "Más detalle técnico: docs/MOBILE_BUILDS.md §5 · docs/EAS_CREDENTIALS_HANDOFF_7_5.md")

    add_h1(doc, "¿Para qué sirven las invitaciones a GitHub?")
    add_para(doc, "Invitaciones a estos dos repositorios:")
    add_bullets(
        doc,
        [
            "TMS Tigerhawk V2 — TMS (panel web / servidor en Vercel): "
            "https://github.com/AI-Developer-Smarter-Business/TMS_tigerhawk_V2",
            "Tigerhawk Mobile App — app móvil (Expo): "
            "https://github.com/AI-Developer-Smarter-Business/tigerhawk_mobile_app",
        ],
    )
    add_para(
        doc,
        "No son para programar obligatoriamente, sino para que el cliente y el equipo de "
        "desarrollo trabajen sobre la misma copia del proyecto.",
    )
    add_para(doc, "Piénselo como una carpeta compartida en la nube con historial de versiones:")
    add_bullets(
        doc,
        [
            "Si el equipo de desarrollo sube un arreglo o una función nueva, el cliente puede bajarlo (git pull) y tendrá exactamente lo mismo.",
            "Si el cliente (o un asistente de IA en su máquina) hace un cambio y lo sube, el equipo lo ve y continúa desde ahí.",
            "Así no hay dos versiones distintas (la del desarrollador vs la de la oficina) que luego no coinciden en producción.",
        ],
    )
    add_para(
        doc,
        "Eso evita sorpresas del tipo: «en mi entorno funciona X pero en el TMS desplegado no está» "
        "o «la app del conductor no tiene lo que probamos ayer».",
    )
    add_para(doc, "En la práctica:", bold=True)
    add_numbered(
        doc,
        [
            "Recibir la invitación por correo → aceptar en GitHub.",
            "Clonar o actualizar el repo cuando se necesite el último código.",
            "Antes de hacer cambios: traer lo último (git pull), luego trabajar, luego subir (si el cliente publica cambios).",
        ],
    )
    p = doc.add_paragraph()
    run = p.add_run(
        "Nota importante: Antes de subir cambios al repositorio (git push), conviene siempre "
        "ejecutar git pull. Eso descarga primero lo que haya en GitHub y fusiona con lo local. "
        "Si no hace pull antes de push, pueden pisar trabajo ajeno o generar conflictos difíciles "
        "de arreglar. Es el equivalente a abrir el documento más reciente antes de editarlo.",
    )
    run.bold = True
    add_para(
        doc,
        "Resumen: las invitaciones a GitHub son para que TMS y app móvil vivan en un solo lugar "
        "de verdad; cliente y desarrollador ven los mismos archivos y el mismo historial, y se "
        "reduce el riesgo de incompatibilidades.",
    )
    add_para(doc, "Guía técnica breve: docs/GitHub_Setup_Guide.md")

    doc.save(OUT_ES)
    print(f"Wrote {OUT_ES}")


if __name__ == "__main__":
    build_en()
    build_es()
