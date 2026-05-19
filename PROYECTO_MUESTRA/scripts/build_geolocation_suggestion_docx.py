#!/usr/bin/env python3
"""Convert docs/SUGGESTION_FOR_RESOLVING_GEOLOCATION.md to .docx (requires python-docx)."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent.parent
MD_PATH = ROOT / "docs" / "SUGGESTION_FOR_RESOLVING_GEOLOCATION.md"
OUT_PATH = ROOT / "docs" / "SUGGESTION_FOR_RESOLVING_GEOLOCATION.docx"


def strip_inline_md(s: str) -> str:
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"`([^`]+)`", r"\1", s)
    s = s.replace("→", "->")
    return s


def is_table_separator(line: str) -> bool:
    inner = line.strip().strip("|")
    if not inner:
        return False
    for part in inner.split("|"):
        p = part.strip()
        if not p:
            return False
        if not all(c in "-: " for c in p):
            return False
        if "-" not in p:
            return False
    return True


def parse_table_row(line: str) -> list[str]:
    parts = line.strip().split("|")
    if parts and parts[0] == "":
        parts = parts[1:]
    if parts and parts[-1] == "":
        parts = parts[:-1]
    return [strip_inline_md(c.strip()) for c in parts]


def main() -> None:
    if not MD_PATH.is_file():
        print("Missing:", MD_PATH, file=sys.stderr)
        sys.exit(1)

    doc = Document()
    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i].rstrip()

        if not line or line.strip() == "---":
            i += 1
            continue

        if line.startswith("# ") and not line.startswith("## "):
            doc.add_heading(strip_inline_md(line[2:]), level=0)
            i += 1
            continue

        if line.startswith("## "):
            doc.add_heading(strip_inline_md(line[3:]), level=1)
            i += 1
            continue

        if line.strip().startswith("|"):
            rows: list[list[str]] = []
            while i < n and lines[i].strip().startswith("|"):
                row_line = lines[i].rstrip()
                if is_table_separator(row_line):
                    i += 1
                    continue
                rows.append(parse_table_row(row_line))
                i += 1
            if rows:
                ncol = max(len(r) for r in rows)
                for r in rows:
                    while len(r) < ncol:
                        r.append("")
                table = doc.add_table(rows=len(rows), cols=ncol)
                table.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci, text in enumerate(row):
                        table.rows[ri].cells[ci].text = text
            continue

        if re.match(r"^\d+\.\s", line):
            text = strip_inline_md(re.sub(r"^\d+\.\s*", "", line))
            doc.add_paragraph(text, style="List Number")
            i += 1
            continue

        if line.startswith("- "):
            doc.add_paragraph(strip_inline_md(line[2:]), style="List Bullet")
            i += 1
            continue

        doc.add_paragraph(strip_inline_md(line))
        i += 1

    doc.save(OUT_PATH)
    print("Wrote", OUT_PATH)


if __name__ == "__main__":
    main()
